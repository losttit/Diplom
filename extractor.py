import os
import docx2txt
from CTkMessagebox import CTkMessagebox
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Inches
from fpdf import FPDF


def extract_text(filename):
    # Извлечение данных из файлов формата PDF/docx
    if filename.endswith('.pdf'):
        # Создаем объект PdfReader для чтения файла PDF
        pdf_reader = PdfReader(filename)
        text = ''
        # Перебираем все страницы PDF файла и извлекаем текст
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    elif filename.endswith('.docx'):
        # Извлекаем текст из файла DOCX с помощью библиотеки docx2txt
        return docx2txt.process(filename)
    else:
        # Если файл не является ни PDF, ни DOCX, выводим ошибку
        CTkMessagebox(title="Ошибка", message="Выберите файл docx или pdf", icon="cancel")
        print("Ошибка: Выберите файл docx или pdf")
        return ''


def add_docx(text, doc_name, image_exists=True):
    # Сохранение лекции в формате DOCX
    if os.path.exists(doc_name):
        # Если файл с таким именем уже существует, удаляем его
        os.remove(doc_name)
    if not doc_name.endswith('.docx'):
        # Если имя файла не заканчивается на '.docx', добавляем это расширение
        doc_name += '.docx'
    # Создаем новый документ DOCX
    doc = Document()
    # Добавляем изображение в документ, если оно существует
    if image_exists:
        doc.add_picture("output/lecture.jpg", width=Inches(6), height=Inches(6))
    # Добавляем текст в документ
    doc.add_paragraph(text)
    # Сохраняем документ
    doc.save("output/" + doc_name)


def add_pdf(text, pdf_name, image_exists=True):
    # Сохранение лекции в формате PDF
    if os.path.exists(pdf_name):
        # Если файл с таким именем уже существует, удаляем его
        os.remove(pdf_name)
    if not pdf_name.endswith('.pdf'):
        # Если имя файла не заканчивается на '.pdf', добавляем это расширение
        pdf_name += '.pdf'
    # Создаем новый PDF документ
    pdf = FPDF()
    # Добавляем страницу в документ
    pdf.add_page()
    # Устанавливаем шрифт и размер текста
    pdf.add_font("DejaVuSans", "", "resources/DejaVuSans.ttf", uni=True)
    pdf.set_font("DejaVuSans", "", 12)
    # Добавляем изображение в документ, если оно существует
    if image_exists:
        pdf.image("output/lecture.jpg", x=10, y=10, w=190, h=190)
        pdf.ln(200)
    # Добавляем текст в документ
    pdf.multi_cell(0, 10, txt=text, align='L')
    # Сохраняем документ
    pdf.output("output/" + pdf_name)
